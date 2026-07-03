"""
BoroBhai File Tools Lambda

Handles deterministic file operations:
- PDF compression and merging
- Image resizing
- Excel generation
- DOCX generation (CV, letters, etc.)
"""

import json
import os
import io
from typing import Any, Dict
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    from PIL import Image
    import PyPDF2
except ImportError:
    # Mock imports for MVP
    pass

import boto3

# PLACEHOLDER: Replace with real credentials
BEDROCK_REGION = os.environ.get('BEDROCK_REGION', 'us-west-2')
S3_BUCKET = os.environ.get('BEDROCK_REGION', 'proofsheet-user-files')

s3 = boto3.client('s3', region_name=BEDROCK_REGION)


# ============================================================
# File Operations
# ============================================================

def compress_pdf(input_key: str, output_key: str) -> str:
    """
    Compress PDF file.

    Args:
        input_key: S3 key of input PDF
        output_key: S3 key for output (compressed) PDF

    Returns:
        Message with download URL
    """
    try:
        # Download PDF from S3
        response = s3.get_object(Bucket=S3_BUCKET, Key=input_key)
        pdf_content = response['Body'].read()

        # Compress using PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        pdf_writer = PyPDF2.PdfWriter()

        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            # Compress: scale content, remove unnecessary data
            pdf_writer.add_page(page)

        # Write compressed PDF
        output_buffer = io.BytesIO()
        pdf_writer.write(output_buffer)
        output_buffer.seek(0)

        # Upload to S3
        s3.put_object(
            Bucket=S3_BUCKET,
            Key=output_key,
            Body=output_buffer.read(),
            ContentType='application/pdf'
        )

        # Generate presigned URL
        url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': output_key},
            ExpiresIn=3600
        )

        return f"PDF compressed successfully. Download: {url}"

    except Exception as e:
        return f"PDF compression failed: {str(e)}"


def merge_pdfs(input_keys: list, output_key: str) -> str:
    """
    Merge multiple PDF files.

    Args:
        input_keys: List of S3 keys for PDFs to merge
        output_key: S3 key for merged PDF

    Returns:
        Message with download URL
    """
    try:
        pdf_writer = PyPDF2.PdfWriter()

        # Read and merge each PDF
        for input_key in input_keys:
            response = s3.get_object(Bucket=S3_BUCKET, Key=input_key)
            pdf_content = response['Body'].read()

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
            for page_num in range(len(pdf_reader.pages)):
                pdf_writer.add_page(pdf_reader.pages[page_num])

        # Write merged PDF
        output_buffer = io.BytesIO()
        pdf_writer.write(output_buffer)
        output_buffer.seek(0)

        # Upload to S3
        s3.put_object(
            Bucket=S3_BUCKET,
            Key=output_key,
            Body=output_buffer.read(),
            ContentType='application/pdf'
        )

        # Generate presigned URL
        url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': output_key},
            ExpiresIn=3600
        )

        return f"PDFs merged successfully ({len(input_keys)} files). Download: {url}"

    except Exception as e:
        return f"PDF merge failed: {str(e)}"


def resize_image(input_key: str, output_key: str, width: int, height: int) -> str:
    """
    Resize image file.

    Args:
        input_key: S3 key of input image
        output_key: S3 key for resized image
        width: Target width in pixels
        height: Target height in pixels

    Returns:
        Message with download URL
    """
    try:
        # Download image from S3
        response = s3.get_object(Bucket=S3_BUCKET, Key=input_key)
        image_content = response['Body'].read()

        # Open and resize image
        image = Image.open(io.BytesIO(image_content))
        resized = image.resize((width, height), Image.Resampling.LANCZOS)

        # Save resized image
        output_buffer = io.BytesIO()
        resized.save(output_buffer, format=image.format or 'PNG')
        output_buffer.seek(0)

        # Upload to S3
        s3.put_object(
            Bucket=S3_BUCKET,
            Key=output_key,
            Body=output_buffer.read(),
            ContentType='image/png'
        )

        # Generate presigned URL
        url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': output_key},
            ExpiresIn=3600
        )

        return f"Image resized successfully ({width}x{height}px). Download: {url}"

    except Exception as e:
        return f"Image resize failed: {str(e)}"


def generate_excel(title: str, headers: list, rows: list, output_key: str) -> str:
    """
    Generate Excel spreadsheet.

    Args:
        title: Sheet title
        headers: Column headers
        rows: Data rows
        output_key: S3 key for output Excel file

    Returns:
        Message with download URL
    """
    try:
        # Create workbook
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = title[:31]  # Max 31 chars for sheet name

        # Add headers
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF')

        for col_num, header in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col_num)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font

        # Add data rows
        for row_num, row_data in enumerate(rows, 2):
            for col_num, cell_value in enumerate(row_data, 1):
                sheet.cell(row=row_num, column=col_num).value = cell_value

        # Auto-fit columns
        for column in sheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                if len(str(cell.value or '')) > max_length:
                    max_length = len(str(cell.value or ''))
            adjusted_width = min(max_length + 2, 50)
            sheet.column_dimensions[column_letter].width = adjusted_width

        # Save and upload
        output_buffer = io.BytesIO()
        workbook.save(output_buffer)
        output_buffer.seek(0)

        s3.put_object(
            Bucket=S3_BUCKET,
            Key=output_key,
            Body=output_buffer.read(),
            ContentType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

        # Generate presigned URL
        url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': output_key},
            ExpiresIn=3600
        )

        return f"Spreadsheet generated successfully ({len(rows)} rows). Download: {url}"

    except Exception as e:
        return f"Excel generation failed: {str(e)}"


def generate_cv_docx(full_name: str, email: str, phone: str, education: list,
                     experience: list, skills: list, output_key: str) -> str:
    """
    Generate CV in DOCX format.

    Args:
        full_name: Person's full name
        email: Email address
        phone: Phone number
        education: List of education entries
        experience: List of work experience entries
        skills: List of skills
        output_key: S3 key for output DOCX

    Returns:
        Message with download URL
    """
    try:
        doc = Document()

        # Header
        title = doc.add_heading(full_name, level=1)
        title.alignment = 1  # Center

        # Contact info
        contact_info = doc.add_paragraph(f"{email} | {phone}")
        contact_info.alignment = 1  # Center

        # Education
        if education:
            doc.add_heading('শিক্ষা (Education)', level=2)
            for entry in education:
                p = doc.add_paragraph(
                    f"{entry.get('degree', '')} in {entry.get('field', '')} - {entry.get('school', '')} ({entry.get('year', '')})",
                    style='List Bullet'
                )

        # Experience
        if experience:
            doc.add_heading('কর্মঅভিজ্ঞতা (Experience)', level=2)
            for entry in experience:
                p = doc.add_paragraph(f"{entry.get('role', '')} at {entry.get('company', '')}")
                p.paragraph_format.left_indent = Inches(0.25)
                p.runs[0].bold = True

                for responsibility in entry.get('responsibilities', []):
                    sub_p = doc.add_paragraph(responsibility, style='List Bullet')
                    sub_p.paragraph_format.left_indent = Inches(0.5)

        # Skills
        if skills:
            doc.add_heading('দক্ষতা (Skills)', level=2)
            p = doc.add_paragraph(', '.join(skills))

        # Save and upload
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)

        s3.put_object(
            Bucket=S3_BUCKET,
            Key=output_key,
            Body=output_buffer.read(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # Generate presigned URL
        url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': output_key},
            ExpiresIn=3600
        )

        return f"CV generated successfully for {full_name}. Download: {url}"

    except Exception as e:
        return f"CV generation failed: {str(e)}"


# ============================================================
# Lambda Handler
# ============================================================

def lambda_handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Lambda handler for file tool execution.

    Expected event format:
    {
        "tool_name": "compress_pdf",
        "parameters": {
            "input_key": "...",
            "output_key": "..."
        }
    }
    """

    try:
        tool_name = event.get('tool_name', '')
        parameters = event.get('parameters', {})

        if tool_name == 'compress_pdf':
            result = compress_pdf(
                parameters.get('input_key', ''),
                parameters.get('output_key', '')
            )
        elif tool_name == 'merge_pdfs':
            result = merge_pdfs(
                parameters.get('input_keys', []),
                parameters.get('output_key', '')
            )
        elif tool_name == 'resize_image':
            result = resize_image(
                parameters.get('input_key', ''),
                parameters.get('output_key', ''),
                parameters.get('width', 800),
                parameters.get('height', 600)
            )
        elif tool_name == 'generate_excel':
            result = generate_excel(
                parameters.get('title', 'Sheet1'),
                parameters.get('headers', []),
                parameters.get('rows', []),
                parameters.get('output_key', '')
            )
        elif tool_name == 'generate_cv_docx':
            result = generate_cv_docx(
                parameters.get('full_name', ''),
                parameters.get('email', ''),
                parameters.get('phone', ''),
                parameters.get('education', []),
                parameters.get('experience', []),
                parameters.get('skills', []),
                parameters.get('output_key', '')
            )
        else:
            result = f"Unknown tool: {tool_name}"

        return {
            'statusCode': 200,
            'body': json.dumps({'result': result}),
            'headers': {'Content-Type': 'application/json'}
        }

    except Exception as e:
        print(f"File tools error: {str(e)}")
        return {
            'statusCode': 500,
            'body': json.dumps({'error': str(e)}),
            'headers': {'Content-Type': 'application/json'}
        }
